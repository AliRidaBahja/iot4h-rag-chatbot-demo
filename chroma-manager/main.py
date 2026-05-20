# %%
import os
import argparse
from openai import OpenAI
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from docx import Document as DocxDocument
from langchain_core.documents import Document
from dotenv import load_dotenv
from chromadb import HttpClient
from langchain_text_splitters import RecursiveCharacterTextSplitter
import json
import logging
import re 
import pandas as pd
import numpy as np 
from similarity import custom_similarity_search
load_dotenv("../.env")                                                                                                                                                  # !! Remove in docker

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
CHROMA_PORT = 8000  
CHROMA_HOST = "localhost"
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") 
COLLECTION_NAME = os.getenv("COLLECTION_NAME")
EMBEDDING_MODEL= os.getenv("EMBEDDING_MODEL") 
DOCS_FOLDER = "./docs"
client = HttpClient(host=CHROMA_HOST, port=CHROMA_PORT)
embedding = OpenAIEmbeddings(model=EMBEDDING_MODEL) 

# Initialize the vector store with the HTTP client
vector_store = Chroma(
    client=client,
    collection_name=COLLECTION_NAME,
    embedding_function=embedding,
)

def load_docx_tables_portal(path):
    doc = DocxDocument(path)
    documents = []

    for table_index, table in enumerate(doc.tables):
        try:
            rows = table.rows
            if len(rows) < 3:
                continue  # not a valid table

            # Row 0, first cell: Title
            title = rows[0].cells[0].text.strip()

            # Row 1: Description and Links
            description = rows[1].cells[0].text.strip()
            links_text = rows[1].cells[1].text
            links = re.findall(r'https?://\S+', links_text)

            # Row 2: Gewerke
            gewerke_text = rows[2].cells[1].text.strip()
            gewerke = re.split(r'[,/\t\n]+', gewerke_text)
            gewerke = [g.strip() for g in gewerke if g.strip()]

            metadata = {
                "source": path,
                "table_index": table_index,
                "title": title,
                "link": ", ".join(links),          # convert list to string
                "gewerke": ", ".join(gewerke)      # convert list to string          
                }

            enriched_content = f"""{metadata["title"]}

Gewerke: {metadata["gewerke"]}
Links: {metadata["link"]:}

{description}
"""
            documents.append(
                Document(
                    page_content=enriched_content,
                    metadata=metadata
                )
            )

        except Exception as e:
            print(f"[WARN] Failed to parse table {table_index}: {e}")
            continue

    return documents

def load_docx_tables(path):
    doc = DocxDocument(path)
    documents = []

    for table_index, table in enumerate(doc.tables):
        try:
            rows = table.rows
            if len(rows) < 3:
                continue  # not a valid table


            title = rows[0].cells[0].text.strip()

            description = rows[1].cells[0].text.strip()
            sensors_text = rows[2].cells[0].text.strip()
            sensors = re.split(r'[,/\t\n]+', sensors_text)
            sensors_list = [s.strip() for s in sensors if s.strip()]
            links = rows[3].cells[0].text.strip()
            links_list = re.findall(r'https?://\S+', links)
            gewerke_text = rows[4].cells[1].text.strip()
            gewerke = re.split(r'[,/\t\n]+', gewerke_text)
            gewerke_list = [g.strip() for g in gewerke if g.strip()]
            status = rows[5].cells[1].text.strip()
            nutzungsphase = rows[6].cells[1].text.strip()

            metadata = {
                "source": path,
                "table_index": table_index,
                "title": title,
                "sensors":  ", ".join(sensors_list),
                "link": ", ".join(links_list),          # convert list to string
                "gewerke": ", ".join(gewerke_list),      # convert list to string  
                "status": status,
                "nutzungsphase": nutzungsphase
                }

            enriched_content =  f"""Title: {metadata["title"]}

Sensors: {metadata["sensors"]}
Links: {metadata["link"]}
Gewerke: {metadata["gewerke"]}
Status: {metadata["status"]}
Nutzungsphase: {metadata["nutzungsphase"]}
{description}
"""
            documents.append(
                Document(
                    page_content=enriched_content,
                    metadata=metadata
                )
            )

        except Exception as e:
            print(f"[WARN] Failed to parse table {table_index}: {e}")
            continue

    return documents

def add_documents(file_path):
    file_name = os.path.basename(file_path)

    if not file_name.endswith(".docx"):
        logging.info(f"Skipped: {file_name} is not a .docx file.")
        return

    try:
        doc = load_docx_tables(file_path)
        vector_store.add_documents(doc)
        logging.info(f"✅ Successfully indexed: {file_name}")
    except Exception as e:
        logging.error(f"❌ Failed to index {file_name}: {e}")


def add_manual_document(
    title: str,
    description: str,
    links: list,
    gewerke: list,
    sensors: list = None,
    status: str = "",
    nutzungsphase: str = "",
    source: str = "manual_entry",
    table_index: int = -1,
):
    try:
        metadata = {
            "source": source,
            "table_index": table_index,
            "title": title,
            "link": ", ".join(links),
            "gewerke": ", ".join(gewerke),
            "sensors": ", ".join(sensors) if sensors else "",
            "status": status,
            "nutzungsphase": nutzungsphase,
        }

        enriched_content = f"""Title: {metadata["title"]}
Sensors: {metadata["sensors"]}
Gewerke: {metadata["gewerke"]}
Links: {metadata["link"]}
Status: {metadata["status"]}
Nutzungsphase: {metadata["nutzungsphase"]}

{description}
"""

        doc = Document(page_content=enriched_content, metadata=metadata)
        vector_store.add_documents([doc])
        logging.info(f"✅ Manually added document: {title}")

    except Exception as e:
        logging.error(f"❌ Failed to manually add document '{title}': {e}")

def add_manual_document_interactively():
    print("\n📄 Manuelle Dokumenteingabe starten...\n")

    title = input("Titel: ").strip()
    description = input("Beschreibung: ").strip()

    links_raw = input("Links (kommagetrennt, optional): ").strip()
    links = [link.strip() for link in links_raw.split(",") if link.strip()]

    gewerke_raw = input("Gewerke (kommagetrennt, optional): ").strip()
    gewerke = [g.strip() for g in gewerke_raw.split(",") if g.strip()]

    sensors_raw = input("Sensoren (kommagetrennt, optional): ").strip()
    sensors = [s.strip() for s in sensors_raw.split(",") if s.strip()]

    status = input("Status (optional): ").strip()
    nutzungsphase = input("Nutzungsphase (optional): ").strip()

    add_manual_document(
        title=title,
        description=description,
        links=links,
        gewerke=gewerke,
        sensors=sensors,
        status=status,
        nutzungsphase=nutzungsphase
    )

DISPLAY_ORDER = [
    "source",
    "table_index",
    "title",
    "sensors",
    "link",
    "gewerke",
    "status",
    "nutzungsphase"
]

def list_documents():
    try:
        results = vector_store.get()
        documents = results.get("documents", [])
        metadatas = results.get("metadatas", [])

        print(f"\n📂 Dokumente in der Sammlung '{COLLECTION_NAME}':\n")

        for i, (metadata, content) in enumerate(zip(metadatas, documents)):
            print(f"[{i}]")
            for key in DISPLAY_ORDER:
                if key in metadata:
                    print(f"  {key.capitalize():<14}: {metadata[key]}")
            print(f"  Content       : {content[:400]}...")  # optional truncation
            print("-" * 60)

    except Exception as e:
        logging.error(f"Fehler beim Auflisten der Dokumente: {e}")


def delete_document_by_title(pattern: str, regex: bool = False):
    try:
        results = vector_store.get()
        ids_to_delete = []

        for doc_id, metadata in zip(results["ids"], results["metadatas"]):
            title = metadata.get("title", "")
            if regex:
                try:
                    if re.search(pattern, title):
                        ids_to_delete.append(doc_id)
                except re.error as e:
                    logging.error(f"Invalid regex pattern: {pattern}. Error: {e}")
                    return
            else:
                if pattern.lower() in title.lower():
                    ids_to_delete.append(doc_id)

        if not ids_to_delete:
            logging.warning(f"No documents found matching title pattern: '{pattern}'")
            return

        vector_store.delete(ids=ids_to_delete)
        logging.info(f"Deleted {len(ids_to_delete)} document(s) with title: '{pattern}'")

    except Exception as e:
        logging.error(f"Error deleting documents by title: {e}")

def delete_document_by_source(source_path: str):
    try:
        results = vector_store.get()
        ids_to_delete = [
            doc_id for doc_id, metadata in zip(results["ids"], results["metadatas"])
            if metadata.get("source") == source_path
        ]
        if not ids_to_delete:
            logging.warning(f"No documents found for source: {source_path}")
            return

        vector_store.delete(ids=ids_to_delete)
        logging.info(f"Deleted {len(ids_to_delete)} document(s) from source: {source_path}")
    except Exception as e:
        logging.error(f"Error deleting document with source {source_path}: {e}")

def interactive_query_loop():
    print("\n🔎 Interaktiver Modus gestartet.")
    try:
        k = int(input("🔢 Wie viele Ergebnisse möchtest du sehen? (Standard: 3) > ").strip() or "3")
    except ValueError:
        print("⚠️  Ungültige Eingabe. Standardwert 3 wird verwendet.")
        k = 3

    print("\n💬 Gib eine Frage ein (oder 'exit' zum Beenden):")
    while True:
        query = input("Frage > ").strip()
        if query.lower() in {"exit", "quit"}:
            print("👋 Auf Wiedersehen!")
            break

        try:
            results = custom_similarity_search(query, k=k, threshold=0.7)
            if not results:
                print("🚫 Keine relevanten Ergebnisse gefunden.")
                continue

            print(f"\n📚 Gefundene Anwendungsfälle:")
            for i, doc in enumerate(results):
                meta = doc.metadata
                print(f"\n[{i+1}] {meta.get('title')}")
                print(f"     Sensors        : {meta.get('sensors', '-')}")
                print(f"     Gewerke        : {meta.get('gewerke', '-')}")
                print(f"     Status         : {meta.get('status', '-')}")
                print(f"     Nutzungsphase  : {meta.get('nutzungsphase', '-')}")
                print(f"     Links          : {meta.get('link', '-')}")
                print(f"     Quelle         : {meta.get('source', '-')}")
                print(f"     Inhalt         : {doc.page_content[:400]}...")  # truncate for readability

        except Exception as e:
            print(f"❌ Fehler bei der Abfrage: {e}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Manage Chroma Vector Store Documents")
    parser.add_argument("--add", type=str, help="Path to .docx file to add")
    parser.add_argument("--list", action="store_true", help="List all documents in the vector store")
    parser.add_argument("--delete", type=str, help="Delete documents by exact source path")
    parser.add_argument("--delete-title", type=str, help="Delete documents by (partial) title match")
    parser.add_argument("--delete-title-regex", action="store_true", help="Interpret title pattern as regex")
    parser.add_argument("--add-manual-interactive", action="store_true", help="Manually add a document via prompts")
    parser.add_argument("--query", action="store_true", help="Start interactive query loop")
    args = parser.parse_args()

    if args.add:
        add_documents(args.add)
    elif args.list:
        list_documents()
    elif args.delete:
        delete_document_by_source(args.delete)
    elif args.delete_title:
        delete_document_by_title(args.delete_title, regex=args.delete_title_regex)
    elif args.add_manual_interactive:
        add_manual_document_interactively()
    elif args.query:
        interactive_query_loop()
    else:
        parser.print_help()

